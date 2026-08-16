"""
Extraction du texte et des champs structurés depuis les documents sources.

Deux niveaux d'extraction :

1. **Texte** : extraction du contenu brut d'un document, quel que soit le format :
   - `.docx` via python-docx (paragraphes + tables),
   - `.pdf` via PyMuPDF.
   Le texte est normalisé (lignes propres, pas de doublons de sauts de ligne).

2. **Champs structurés** : repérage par règles (regex/NLP léger) des informations clés
   d'un dossier d'appel d'offres pour les couches aval (résumé, checklist d'éligibilité,
   RAG). Pour les dossiers-types ARCOP, les montants/dates sont des *placeholders*
   ex. « [Insérer le montant en FCFA] » : l'extraction documente ce qu'elle trouve
   (valeur réelle OU placeholder), avec la source qui a permis de la repérer.

Usage :
    python src/extraction.py --doc path/to/file.docx --db data/processed/extraction.db
    python src/extraction.py --dir data/raw/dossiers_types --db data/processed/extraction.db
"""

from __future__ import annotations

import argparse
import re
import sqlite3
from dataclasses import dataclass, asdict
from datetime import datetime, timezone
from pathlib import Path

# --- Extraction du texte brut -------------------------------------------------


def extract_text(path: Path) -> str:
    """Extrait le texte d'un document .docx ou .pdf selon son extension."""
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _extract_text_docx(path)
    if suffix == ".pdf":
        return _extract_text_pdf(path)
    raise ValueError(f"Format non supporté : {suffix} (fichier : {path.name})")


def _extract_text_docx(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts: list[str] = []

    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))

    return _normalize_text("\n".join(parts))


def _extract_text_pdf(path: Path) -> str:
    import pymupdf

    parts: list[str] = []
    with pymupdf.open(str(path)) as document:
        for page in document:
            text = page.get_text()
            if text.strip():
                parts.append(text)
    return _normalize_text("\n".join(parts))


def _normalize_text(text: str) -> str:
    """Normalise le texte extrait : lignes sans espaces parasites, pas de triple sauts."""
    lines = [line.strip() for line in text.splitlines()]
    lines = [line for line in lines if line]
    return "\n".join(lines)


# --- Extraction des champs structurés --------------------------------------

@dataclass
class ChampsExtraits:
    objet: str | None = None
    montant_previsionnel: str | None = None
    garantie_soumission: str | None = None
    delai_execution: str | None = None
    validite_offres: str | None = None
    date_limite_depot: str | None = None
    lieu_depot: str | None = None
    contact_consultation: str | None = None


# Un "champ" = (label, regex à trouver dans le texte, valeur capturée).
# ordre des patterns : le plus spécifique d'abord.
#
# NB : les dossiers-types ARCOP sont des modèles : les dates/montants réels y sont des
# placeholders ex. "[Insérer la date et l'heure]". Les patterns ci-dessous capturent
# donc volontairement les deux formes — valeur réelle OU placeholder — et le champ
# is_placeholder permet de distinguer les deux en aval.
DATE_LITERALE_PATTERN = (
    r"\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}"        # 10/07/2026 ou 10-07-2026
    r"|\d{1,2}\s+(?:janvier|février|fevrier|mars|avril|mai|juin|juillet|août|aout|"
    r"septembre|octobre|novembre|décembre|decembre)\s+\d{4}"   # 10 juillet 2026
    r"|\[[^\]]*(?:date|heure|Insérer la date)[^\]]*\]"          # [Insérer la date et l'heure]
)

FIELD_PATTERNS: list[tuple[str, str]] = [
    # objet
    ("objet", r"(?:^|\n)\s*Objet\s*[:\-]?\s*(.+)$"),
    # montant prévisionnel
    ("montant_previsionnel", r"Le montant prévisionnel des travaux est de\s*(.+?)\."),
    ("montant_previsionnel", r"Le montant prévisionnel du marché est de\s*(.+?)\."),
    # garantie de soumission
    ("garantie_soumission", r"garantie de soumission[^.]*?d'un montant de\s*(.+?)[.;]"),
    ("garantie_soumission", r"garantie de soumission[^.]*?montant de\s*(.+?)[.;]"),
    # délai d'exécution
    ("delai_execution", r"Le délai d'exécution est de\s*(.+?)[.;]"),
    ("delai_execution", r"délai d'exécution\s*(?:est|sera) de\s*(.+?)[.;]"),
    # validité des offres
    ("validite_offres", r"engagés par leur offre pendant une période de\s*(.+?)\s*à compter"),
    ("validite_offres", r"valides pendant\s*(.+?)\s*(?:à compter|à partir|depuis)"),
    # date limite / réception des offres : on veut la date, pas le mot "paragraphe".
    ("date_limite_depot",
     r"(?:date limite (?:du dépôt|de dépôt|de soumission|fixée|de remise)|"
     r"réception des offres|soumises.{0,40}au plus tard le)\s*"
     r"(?P<date>" + DATE_LITERALE_PATTERN + r")"),
    # lieu de dépôt
    ("lieu_depot", r"déposées\s*au\s*(.+?)[.;\n]"),
    ("lieu_depot", r"(?:adressées|réceptionnées)\s*(?:au|à)\s*(.+?)[.;\n]"),
    # contact consultation
    ("contact_consultation", r"retirer à titre gratuit|consultés gratuitement\s*(?:au|à|chez)\s*(.+?)[.;\n]"),
    ("contact_consultation", r"consulté gratuitement\s*(?:au|à|chez)\s*(.+?)[.;\n]"),
]

# Pour les dossiers-types, un placeholder est une valeur VALIDE à documenter,
# pas un échec : on le garde tel quel.
PLACEHOLDER_PATTERN = re.compile(r"\[[^\]]+\]|\bV\b|\bN\b|\bX\b")


def is_placeholder(value: str | None) -> bool:
    """True si la valeur extraite est un placeholder de modèle (ex. '[Insérer ...]')."""
    if not value:
        return False
    return bool(PLACEHOLDER_PATTERN.search(value)) or "insérer" in value.lower()


def extract_fields(text: str) -> ChampsExtraits:
    """Extrait les champs structurés d'un texte de DAO/DPAO par règles."""
    champs = ChampsExtraits()
    for label, pattern in FIELD_PATTERNS:
        match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
        if match:
            value = match.group(1).strip()
            if value:
                setattr(champs, label, value)
    return champs


# --- Stockage SQLite ---------------------------------------------------------


def save_document_to_sqlite(*, url: str, local_path: Path, titre: str, categorie: str | None,
                            texte: str, champs: ChampsExtraits, db_path: Path) -> None:
    """Persiste un document (métadonnées + texte + champs) dans la base d'extraction."""
    db_path.parent.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(db_path)
    now = datetime.now(timezone.utc).isoformat()

    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            url TEXT UNIQUE,
            local_path TEXT,
            titre TEXT NOT NULL,
            categorie TEXT,
            scraped_at TEXT NOT NULL,
            texte TEXT NOT NULL
        )
        """
    )
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS champs_extraits (
            document_id INTEGER NOT NULL,
            champ TEXT NOT NULL,
            valeur TEXT,
            is_placeholder INTEGER NOT NULL DEFAULT 0,
            PRIMARY KEY (document_id, champ),
            FOREIGN KEY (document_id) REFERENCES documents (id)
        )
        """
    )

    conn.execute(
        """
        INSERT INTO documents (url, local_path, titre, categorie, scraped_at, texte)
        VALUES (:url, :local_path, :titre, :categorie, :scraped_at, :texte)
        ON CONFLICT(url) DO UPDATE SET
            local_path=excluded.local_path,
            titre=excluded.titre,
            categorie=excluded.categorie,
            scraped_at=excluded.scraped_at,
            texte=excluded.texte
        """,
        {
            "url": url or str(local_path),
            "local_path": str(local_path),
            "titre": titre,
            "categorie": categorie,
            "scraped_at": now,
            "texte": texte,
        },
    )
    document_id = conn.execute(
        "SELECT id FROM documents WHERE url = ?", (url or str(local_path),)
    ).fetchone()[0]

    for champ, valeur in asdict(champs).items():
        if valeur is None:
            continue
        conn.execute(
            """
            INSERT INTO champs_extraits (document_id, champ, valeur, is_placeholder)
            VALUES (:document_id, :champ, :valeur, :is_placeholder)
            ON CONFLICT(document_id, champ) DO UPDATE SET
                valeur=excluded.valeur,
                is_placeholder=excluded.is_placeholder
            """,
            {
                "document_id": document_id,
                "champ": champ,
                "valeur": valeur,
                "is_placeholder": int(is_placeholder(valeur)),
            },
        )

    conn.commit()
    conn.close()


# --- CLI --------------------------------------------------------------------


def _ingest_file(path: Path, db_path: Path, *, url: str | None = None,
                 titre: str | None = None, categorie: str | None = None) -> None:
    texte = extract_text(path)
    champs = extract_fields(texte)
    save_document_to_sqlite(
        url=url or str(path),
        local_path=path,
        titre=titre or path.stem,
        categorie=categorie,
        texte=texte,
        champs=champs,
        db_path=db_path,
    )


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--doc", type=Path, help="Extrait un fichier unique (.docx ou .pdf)"
    )
    parser.add_argument(
        "--dir", type=Path, help="Extrait tous les documents d'un dossier"
    )
    parser.add_argument("--db", type=Path, default=Path("data/processed/extraction.db"))
    args = parser.parse_args()

    if args.doc:
        _ingest_file(args.doc, args.db)
    elif args.dir:
        for path in sorted(args.dir.iterdir()):
            if path.is_file() and path.suffix.lower() in (".docx", ".pdf"):
                _ingest_file(path, args.db)
    else:
        parser.error("Fournir --doc ou --dir")


if __name__ == "__main__":
    main()