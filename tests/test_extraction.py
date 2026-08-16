import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from extraction import (  # noqa: E402
    extract_fields,
    extract_text,
    is_placeholder,
    save_document_to_sqlite,
    ChampsExtraits,
)

SAMPLE_TEXT = (
    "AVIS D'APPEL D'OFFRES N° 001/2026\n"
    "Objet : Réalisation de forages à énergie solaire\n"
    "Le montant prévisionnel des travaux est de 450 000 000 FCFA.\n"
    "Les offres doivent être soumises au plus tard le 10/07/2026.\n"
    "Le délai d'exécution est de 12 mois.\n"
    "Les candidats resteront engagés par leur offre pendant une période de 90 jours à compter de la date limite du dépôt des offres."
)

PLACEHOLDER_TEXT = (
    "Le montant prévisionnel des travaux est de [Insérer le montant prévisionnel du marché].\n"
    "Les offres doivent être soumises au plus tard le [Insérer la date et l'heure].\n"
)


def test_extract_fields_real_values():
    champs = extract_fields(SAMPLE_TEXT)
    assert champs.objet == "Réalisation de forages à énergie solaire"
    assert champs.montant_previsionnel == "450 000 000 FCFA"
    assert champs.date_limite_depot == "10/07/2026"
    assert champs.delai_execution == "12 mois"
    assert champs.validite_offres == "90 jours"


def test_extract_fields_placeholders_are_valid():
    champs = extract_fields(PLACEHOLDER_TEXT)
    # Les placeholders sont des valeurs documentées, pas des échecs d'extraction.
    assert champs.montant_previsionnel == "[Insérer le montant prévisionnel du marché]"
    assert champs.date_limite_depot == "[Insérer la date et l'heure]"


def test_is_placeholder():
    assert is_placeholder("[Insérer le montant en FCFA]") is True
    assert is_placeholder("450 000 000 FCFA") is False
    assert is_placeholder(None) is False
    assert is_placeholder("") is False


def test_extract_text_docx(tmp_path):
    import docx

    doc_path = tmp_path / "test.docx"
    document = docx.Document()
    document.add_paragraph("Premier paragraphe du DAO.")
    document.add_paragraph("Montant prévisionnel : 100 000 FCFA.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Référence"
    table.rows[0].cells[1].text = "AO-2026-00001"
    document.save(str(doc_path))

    texte = extract_text(doc_path)
    assert "Premier paragraphe du DAO." in texte
    assert "Montant prévisionnel : 100 000 FCFA." in texte
    assert "Référence" in texte and "AO-2026-00001" in texte


def test_extract_text_unknown_format(tmp_path):
    bogus = tmp_path / "fichier.txt"
    bogus.write_text("sans importance", encoding="utf-8")
    try:
        extract_text(bogus)
        assert False, "devrait lever ValueError"
    except ValueError:
        pass


def test_save_document_to_sqlite_roundtrip(tmp_path):
    db = tmp_path / "extraction.db"
    doc_path = tmp_path / "dao.docx"
    import docx

    document = docx.Document()
    document.add_paragraph("DAO TEST")
    document.save(str(doc_path))

    save_document_to_sqlite(
        url="https://example.org/dao.docx",
        local_path=doc_path,
        titre="DAO TEST",
        categorie="Dossiers typesTravaux",
        texte="DAO TEST\nObjet : Test",
        champs=ChampsExtraits(objet="Test"),
        db_path=db,
    )

    import sqlite3

    conn = sqlite3.connect(str(db))
    n_docs = conn.execute("SELECT COUNT(*) FROM documents").fetchone()[0]
    n_champs = conn.execute("SELECT COUNT(*) FROM champs_extraits").fetchone()[0]
    conn.close()
    assert n_docs == 1
    assert n_champs == 1


if __name__ == "__main__":
    import tempfile

    t = tempfile.TemporaryDirectory()
    test_extract_fields_real_values()
    test_extract_fields_placeholders_are_valid()
    test_is_placeholder()
    test_extract_text_docx(Path(t.name))
    test_extract_text_unknown_format(Path(t.name))
    test_save_document_to_sqlite_roundtrip(Path(t.name))
    t.cleanup()
    print("Tous les tests locaux passent.")